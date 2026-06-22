import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
import io
import re
from datetime import datetime



# ─── Page Config ──────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="SOW Field Updater",
    page_icon="📄",
    layout="wide"
)



st.title("📄 SOW Field Updater")
st.markdown("Upload your SOW document, edit the key fields, and download the updated file.")



# ─── Helper: Find and replace text in a paragraph ─────────────────────────────
def replace_text_in_paragraph(para, old_text, new_text):
    """Replace text across runs in a paragraph, preserving formatting."""
    full_text = "".join(run.text for run in para.runs)
    if old_text in full_text:
        new_full = full_text.replace(old_text, new_text)
        # Clear all runs and put text in first run
        for i, run in enumerate(para.runs):
            if i == 0:
                run.text = new_full
            else:
                run.text = ""
        return True
    return False



def replace_in_table_cell(cell, old_text, new_text):
    """Replace text in all paragraphs within a table cell."""
    replaced = False
    for para in cell.paragraphs:
        if replace_text_in_paragraph(para, old_text, new_text):
            replaced = True
    return replaced



def replace_in_document(doc, old_text, new_text):
    """Replace text throughout the entire document."""
    count = 0
    # Search paragraphs
    for para in doc.paragraphs:
        if replace_text_in_paragraph(para, old_text, new_text):
            count += 1
    # Search tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if replace_in_table_cell(cell, old_text, new_text):
                    count += 1
                # Nested tables
                for nested_table in cell.tables:
                    for nrow in nested_table.rows:
                        for ncell in nrow.cells:
                            if replace_in_table_cell(ncell, old_text, new_text):
                                count += 1
    return count



def find_text_in_doc(doc, search_text):
    """Find text and return its location for preview."""
    results = []
    for i, para in enumerate(doc.paragraphs):
        if search_text.lower() in para.text.lower():
            results.append(f"Paragraph {i}: ...{para.text[:80]}...")
    for t, table in enumerate(doc.tables):
        for r, row in enumerate(table.rows):
            for c, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    if search_text.lower() in para.text.lower():
                        results.append(f"Table {t}, Row {r}, Col {c}: ...{para.text[:80]}...")
    return results



# ─── Build Milestone Table in DOCX ────────────────────────────────────────────
def build_milestone_table_in_cell(cell, milestones):
    """Replace milestone content in a cell with a formatted table."""
    # Clear existing content
    for para in cell.paragraphs:
        para.clear()



    headers = ["Milestone Name", "Payment Milestone Dt", "Monthly Milestone 85%", "Quality 15%", "Invoice Amount"]



    tbl = OxmlElement('w:tbl')
    tblPr = OxmlElement('w:tblPr')
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)
    tbl.append(tblPr)



    def make_row(texts, bold=False, bg_color=None):
        tr = OxmlElement('w:tr')
        for text in texts:
            tc = OxmlElement('w:tc')
            tcPr = OxmlElement('w:tcPr')
            if bg_color:
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), bg_color)
                shd.set(qn('w:val'), 'clear')
                tcPr.append(shd)
            tc.append(tcPr)
            p = OxmlElement('w:p')
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            if bold:
                b = OxmlElement('w:b')
                rPr.append(b)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '20')
            rPr.append(sz)
            r.append(rPr)
            t = OxmlElement('w:t')
            t.text = str(text)
            r.append(t)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        return tr



    tbl.append(make_row(headers, bold=True, bg_color="DCE6F1"))
    for m in milestones:
        tbl.append(make_row([m["name"], m["date"], m["monthly"], m["quality"], m["invoice"]]))



    cell._tc.append(tbl)



# ─── Session State Init ───────────────────────────────────────────────────────
if "milestones" not in st.session_state:
    st.session_state.milestones = [
        {"name": "Jul", "date": "30th Jul 2026", "monthly": "€ 15953", "quality": "€ 2815", "invoice": "€ 18768"},
        {"name": "Aug", "date": "30th Aug 2026", "monthly": "€ 14566", "quality": "€ 2570", "invoice": "€ 17136"},
        {"name": "Sep", "date": "30th Sep 2026", "monthly": "€ 14566", "quality": "€ 2570", "invoice": "€ 17136"},
    ]



if "original_values" not in st.session_state:
    st.session_state.original_values = {
        "start_date": "01-Jul-2026",
        "end_date": "30-Sep-2026",
        "partner_cost": "€ 53,040.00",
        "release_start": "01-Jul-2026",
        "release_end": "30-Sep-2026",
        "sow_approved_by": "duraipandi.palaniyandi@philips.com",
        "sow_approved_on": "10-Jun-2026",
    }



# ─── File Upload ──────────────────────────────────────────────────────────────
st.markdown("---")
uploaded_file = st.file_uploader("📁 Upload SOW Document (.docx)", type=["docx"])



if uploaded_file:
    st.success(f"✅ Uploaded: **{uploaded_file.name}**")



st.markdown("---")



# ─── Form Fields ──────────────────────────────────────────────────────────────
col1, col2 = st.columns(2)



with col1:
    st.subheader("📅 SOW Dates")
    start_date = st.text_input("Start Date (DD-MMM-YYYY)", value=st.session_state.original_values["start_date"])
    end_date = st.text_input("End Date (DD-MMM-YYYY)", value=st.session_state.original_values["end_date"])



with col2:
    st.subheader("💶 Partner Fixed Cost")
    st.caption("Rate card differs per project")
    partner_cost = st.text_input("Total Cost (Phoenix Rate Card)", value=st.session_state.original_values["partner_cost"])



st.markdown("---")
st.subheader("🏁 Payment Milestones")
st.caption("Edit directly in the table. Add or remove rows as needed.")



# Milestone table header
hcols = st.columns([1.5, 2, 2, 2, 2, 0.5])
for col, h in zip(hcols, ["Month", "Payment Date", "Monthly (85%)", "Quality (15%)", "Invoice Total", "Del"]):
    col.markdown(f"**{h}**")



rows_to_delete = []
for i, m in enumerate(st.session_state.milestones):
    c1, c2, c3, c4, c5, c6 = st.columns([1.5, 2, 2, 2, 2, 0.5])
    st.session_state.milestones[i]["name"]    = c1.text_input("", value=m["name"],    key=f"mn_{i}", label_visibility="collapsed")
    st.session_state.milestones[i]["date"]    = c2.text_input("", value=m["date"],    key=f"md_{i}", label_visibility="collapsed")
    st.session_state.milestones[i]["monthly"] = c3.text_input("", value=m["monthly"], key=f"mm_{i}", label_visibility="collapsed")
    st.session_state.milestones[i]["quality"] = c4.text_input("", value=m["quality"], key=f"mq_{i}", label_visibility="collapsed")
    st.session_state.milestones[i]["invoice"] = c5.text_input("", value=m["invoice"], key=f"mi_{i}", label_visibility="collapsed")
    if c6.button("✕", key=f"del_{i}"):
        rows_to_delete.append(i)



for i in sorted(rows_to_delete, reverse=True):
    st.session_state.milestones.pop(i)
    st.rerun()



if st.button("➕ Add Milestone Row"):
    st.session_state.milestones.append({"name": "", "date": "", "monthly": "", "quality": "", "invoice": ""})
    st.rerun()



st.markdown("---")



col3, col4 = st.columns(2)
with col3:
    st.subheader("🚀 Release Dates")
    release_start = st.text_input("Release Start Date (DD-MMM-YYYY)", value=st.session_state.original_values["release_start"])
    release_end   = st.text_input("End of Warranty Period (DD-MMM-YYYY)", value=st.session_state.original_values["release_end"])



with col4:
    st.subheader("✅ SOW Approval")
    sow_approved_by = st.text_input("SOW Approved By", value=st.session_state.original_values["sow_approved_by"])
    sow_approved_on = st.text_input("SOW Approved On (DD-MMM-YYYY)", value=st.session_state.original_values["sow_approved_on"])



st.markdown("---")



# ─── Generate & Download ──────────────────────────────────────────────────────
st.subheader("⬇ Generate Updated SOW")



if st.button("⚡ Generate Updated Document", type="primary", use_container_width=True):
    if not uploaded_file:
        st.warning("⚠ No file uploaded. Please upload a .docx file first.")
    else:
        with st.spinner("Updating document fields..."):
            try:
                # Load original document
                doc_bytes = io.BytesIO(uploaded_file.read())
                doc = Document(doc_bytes)



                orig = st.session_state.original_values



                # ── Replace each field ──────────────────────────────────────
                replacements = {
                    orig["start_date"]:       start_date,
                    orig["end_date"]:         end_date,
                    orig["partner_cost"]:     partner_cost,
                    orig["release_start"]:    release_start,
                    orig["release_end"]:      release_end,
                    orig["sow_approved_by"]:  sow_approved_by,
                    orig["sow_approved_on"]:  sow_approved_on,
                }



                total_replaced = 0
                for old, new in replacements.items():
                    if old != new and old:
                        n = replace_in_document(doc, old, new)
                        total_replaced += n



                # ── Update milestone table ──────────────────────────────────
                # Find the cell that contains milestone data and rebuild it
                milestone_keywords = ["Monthly Milestone", "Payment Milestone Dt", "Milestone Name"]
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            cell_text = cell.text
                            if any(kw in cell_text for kw in milestone_keywords):
                                build_milestone_table_in_cell(cell, st.session_state.milestones)



                # ── Save to buffer ──────────────────────────────────────────
                output = io.BytesIO()
                doc.save(output)
                output.seek(0)



                orig_name = uploaded_file.name.replace(".docx", "").replace(".doc", "")
                out_name  = f"{orig_name}_Updated.docx"



                st.success(f"✅ Document updated! {total_replaced} field(s) replaced.")



                st.download_button(
                    label="📥 Download Updated SOW",
                    data=output,
                    file_name=out_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )



                # Update session state with new values as "original" for next run
                st.session_state.original_values = {
                    "start_date":      start_date,
                    "end_date":        end_date,
                    "partner_cost":    partner_cost,
                    "release_start":   release_start,
                    "release_end":     release_end,
                    "sow_approved_by": sow_approved_by,
                    "sow_approved_on": sow_approved_on,
                }



            except Exception as e:
                st.error(f"❌ Error processing document: {e}")
                st.exception(e)



st.markdown("---")
st.caption("SOW Field Updater · Cognizant OSCAR Project · Built with Streamlit")